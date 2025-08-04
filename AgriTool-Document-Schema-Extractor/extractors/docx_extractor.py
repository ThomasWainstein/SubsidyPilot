"""
DOCX document extractor for structured form extraction.
"""

import asyncio
import logging
from typing import Dict, Any

from docx import Document


class DOCXExtractor:
    """DOCX document schema extractor."""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
    
    async def extract_schema(self, file_path: str) -> Dict[str, Any]:
        """
        Extract schema from DOCX document.
        
        Args:
            file_path: Path to DOCX file.
            
        Returns:
            Extracted schema data with metadata.
        """
        self.logger.info("Starting DOCX extraction", extra={
            "file_path": file_path,
            "extractor": "DOCXExtractor"
        })
        
        # Run extraction in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_docx_sync, file_path)
    
    def _extract_docx_sync(self, file_path: str) -> Dict[str, Any]:
        """Synchronous DOCX extraction."""
        fields = []
        raw_unclassified = []
        metadata = {
            "method": "docx_extraction",
            "extractor": "DOCXExtractor",
            "has_tables": False,
            "table_count": 0,
            "paragraph_count": 0
        }
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
                    
                    # Check for form fields in paragraph styles
                    if paragraph.style.name in ['Form Field', 'Input Field']:
                        fields.append(self._create_field_from_paragraph(paragraph))
            
            metadata["paragraph_count"] = paragraph_count
            
            # Extract text and fields from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_data = self._extract_table_fields(table)
                fields.extend(table_data['fields'])
                text_content += table_data['text'] + "\n"
            
            metadata["table_count"] = table_count
            metadata["has_tables"] = table_count > 0
            
            # Parse content for additional form fields
            if text_content.strip():
                from utils.field_parser import FieldParser
                parser = FieldParser(self.logger)
                parsed_data = parser.parse_text_for_fields(text_content)
                fields.extend(parsed_data['fields'])
                raw_unclassified.extend(parsed_data['raw_unclassified'])
            
        except Exception as e:
            self.logger.error("DOCX extraction failed", extra={
                "file_path": file_path,
                "error": str(e),
                "error_type": type(e).__name__
            })
            raw_unclassified.append(f"DOCX extraction error: {str(e)}")
        
        return {
            "fields": fields,
            "raw_unclassified": raw_unclassified,
            "metadata": metadata
        }
    
    def _create_field_from_paragraph(self, paragraph) -> Dict[str, Any]:
        """Create field from paragraph with form field style."""
        text = paragraph.text.strip()
        
        from utils.field_parser import FieldParser
        parser = FieldParser(self.logger)
        
        return {
            "name": parser.normalize_field_name(text),
            "label": text,
            "type": parser.infer_field_type(text),
            "required": parser.is_field_required(text),
            "help": "",
            "source": "paragraph_style"
        }
    
    def _extract_table_fields(self, table) -> Dict[str, Any]:
        """Extract fields from table structure."""
        fields = []
        text_content = ""
        
        from utils.field_parser import FieldParser
        parser = FieldParser(self.logger)
        
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_text.append(cell_text)
                text_content += cell_text + " "
                
                # Check if cell looks like a field label
                if cell_text and len(cell_text) > 3 and len(cell_text) < 100:
                    # Look for patterns that suggest this is a field
                    if ':' in cell_text or '?' in cell_text or cell_text.endswith('*'):
                        fields.append({
                            "name": parser.normalize_field_name(cell_text),
                            "label": cell_text,
                            "type": parser.infer_field_type(cell_text),
                            "required": parser.is_field_required(cell_text),
                            "help": "",
                            "source": f"table_cell_{row_idx}_{cell_idx}"
                        })
            
            text_content += "\n"
        
        return {
            "fields": fields,
            "text": text_content
        }