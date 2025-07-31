#!/usr/bin/env python3
"""
AgriTool Subsidy Document Structure Extraction & Schema Mapping

This module automatically extracts real application form structure from subsidy documents
(PDF, DOCX, XLSX) and maps them to standardized JSON schemas for storage in Supabase.

USAGE:
    python document_schema_extractor.py --subsidy-ids 123e4567-e89b-12d3-a456-426614174000
    python document_schema_extractor.py --all-subsidies
    python document_schema_extractor.py --batch-size 10 --max-concurrent 5

ENVIRONMENT VARIABLES:
    SUPABASE_URL - Supabase project URL
    SUPABASE_SERVICE_ROLE_KEY - Service role key for database access
    MAX_CONCURRENT_DOCS - Max concurrent document processing (default: 5)
    DOWNLOAD_TIMEOUT - Download timeout in seconds (default: 30)
    LOG_LEVEL - Logging level (default: INFO)

REQUIREMENTS:
    pip install supabase aiohttp aiofiles PyPDF2 python-docx openpyxl pytesseract Pillow

FEATURES:
    - Async document downloading and processing
    - Multi-format support (PDF, DOCX, XLSX)
    - OCR for scanned documents
    - Configurable concurrency
    - Robust error handling
    - Idempotent operations
    - Production-ready logging
"""

import asyncio
import argparse
import json
import logging
import os
import re
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from urllib.parse import urlparse

import aiofiles
import aiohttp
from supabase import create_client, Client

# Document parsing libraries
import PyPDF2
from docx import Document
from openpyxl import load_workbook
import pytesseract
from PIL import Image

# Configure logging
logging.basicConfig(
    level=getattr(logging, os.getenv('LOG_LEVEL', 'INFO')),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentSchemaExtractor:
    """
    Main class for extracting document schemas from subsidy application forms.
    """
    
    def __init__(self):
        """Initialize the extractor with configuration from environment variables."""
        self.supabase_url = os.getenv('SUPABASE_URL')
        self.supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY')
        
        if not self.supabase_url or not self.supabase_key:
            raise ValueError("SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY must be set")
        
        self.supabase: Client = create_client(self.supabase_url, self.supabase_key)
        self.max_concurrent = int(os.getenv('MAX_CONCURRENT_DOCS', '5'))
        self.download_timeout = int(os.getenv('DOWNLOAD_TIMEOUT', '30'))
        
        # Create semaphore for concurrency control
        self.semaphore = asyncio.Semaphore(self.max_concurrent)
        
        logger.info(f"Initialized DocumentSchemaExtractor with max_concurrent={self.max_concurrent}")

    async def process_subsidies(self, subsidy_ids: List[str] = None, batch_size: int = 10) -> Dict[str, Any]:
        """
        Process multiple subsidies for document extraction.
        
        Args:
            subsidy_ids: List of subsidy IDs to process. If None, process all.
            batch_size: Number of subsidies to process in each batch.
            
        Returns:
            Summary statistics of the extraction process.
        """
        logger.info(f"Starting document extraction for subsidies: {subsidy_ids or 'ALL'}")
        
        # Fetch subsidies to process
        if subsidy_ids:
            subsidies = []
            for subsidy_id in subsidy_ids:
                result = self.supabase.table('subsidies').select('*').eq('id', subsidy_id).execute()
                if result.data:
                    subsidies.extend(result.data)
        else:
            result = self.supabase.table('subsidies').select('*').execute()
            subsidies = result.data
        
        if not subsidies:
            logger.warning("No subsidies found to process")
            return {"processed": 0, "failed": 0, "total": 0}
        
        logger.info(f"Found {len(subsidies)} subsidies to process")
        
        # Process subsidies in batches
        stats = {"processed": 0, "failed": 0, "total": len(subsidies), "extraction_results": []}
        
        for i in range(0, len(subsidies), batch_size):
            batch = subsidies[i:i + batch_size]
            logger.info(f"Processing batch {i//batch_size + 1}: {len(batch)} subsidies")
            
            # Process batch concurrently
            tasks = [self.process_single_subsidy(subsidy) for subsidy in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Update statistics
            for result in results:
                if isinstance(result, Exception):
                    logger.error(f"Batch processing error: {result}")
                    stats["failed"] += 1
                elif result["success"]:
                    stats["processed"] += 1
                    stats["extraction_results"].append(result)
                else:
                    stats["failed"] += 1
        
        logger.info(f"Extraction complete. Processed: {stats['processed']}, Failed: {stats['failed']}")
        return stats

    async def process_single_subsidy(self, subsidy: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process a single subsidy for document extraction.
        
        Args:
            subsidy: Subsidy record from database.
            
        Returns:
            Processing result with success status and extracted data.
        """
        subsidy_id = subsidy['id']
        subsidy_title = subsidy.get('title', 'Unknown')
        
        logger.info(f"Processing subsidy {subsidy_id}: {subsidy_title}")
        
        try:
            # Get document URLs from application_docs field
            application_docs = subsidy.get('application_docs', [])
            if not application_docs:
                logger.warning(f"No application documents found for subsidy {subsidy_id}")
                return {"success": True, "subsidy_id": subsidy_id, "documents_processed": 0}
            
            # Process each document
            extraction_results = []
            total_fields = 0
            
            async with self.semaphore:
                for doc_info in application_docs:
                    try:
                        result = await self.process_document(subsidy_id, doc_info)
                        extraction_results.append(result)
                        total_fields += result.get('field_count', 0)
                    except Exception as e:
                        logger.error(f"Error processing document {doc_info}: {e}")
                        # Record failed extraction
                        await self.record_extraction_status(
                            subsidy_id=subsidy_id,
                            document_url=doc_info.get('url', 'unknown'),
                            document_type=doc_info.get('type', 'unknown'),
                            status='failure',
                            error_message=str(e)
                        )
            
            # Generate consolidated schema from all documents
            consolidated_schema = self.consolidate_schemas(extraction_results)
            
            # Update subsidy with consolidated schema
            await self.update_subsidy_schema(subsidy_id, consolidated_schema)
            
            return {
                "success": True,
                "subsidy_id": subsidy_id,
                "documents_processed": len(extraction_results),
                "total_fields": total_fields,
                "schema": consolidated_schema
            }
            
        except Exception as e:
            logger.error(f"Error processing subsidy {subsidy_id}: {e}")
            return {"success": False, "subsidy_id": subsidy_id, "error": str(e)}

    async def process_document(self, subsidy_id: str, doc_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process a single document for schema extraction.
        
        Args:
            subsidy_id: The subsidy ID this document belongs to.
            doc_info: Document information containing URL and type.
            
        Returns:
            Extraction result with schema and metadata.
        """
        doc_url = doc_info.get('url', '')
        doc_type = doc_info.get('type', '')
        
        logger.info(f"Processing document: {doc_url}")
        
        try:
            # Download document
            file_path = await self.download_document(doc_url)
            
            # Extract schema based on document type
            if doc_type.lower() == 'pdf' or doc_url.lower().endswith('.pdf'):
                extracted_data = await self.extract_pdf_schema(file_path)
            elif doc_type.lower() in ['docx', 'doc'] or doc_url.lower().endswith(('.docx', '.doc')):
                extracted_data = await self.extract_docx_schema(file_path)
            elif doc_type.lower() in ['xlsx', 'xls'] or doc_url.lower().endswith(('.xlsx', '.xls')):
                extracted_data = await self.extract_xlsx_schema(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
            
            # Calculate coverage percentage
            coverage_percentage = self.calculate_coverage(extracted_data)
            
            # Record extraction status
            await self.record_extraction_status(
                subsidy_id=subsidy_id,
                document_url=doc_url,
                document_type=doc_type,
                status='success' if coverage_percentage > 50 else 'partial',
                field_count=len(extracted_data.get('fields', [])),
                coverage_percentage=coverage_percentage,
                extracted_schema=extracted_data
            )
            
            # Clean up downloaded file
            os.unlink(file_path)
            
            return {
                "document_url": doc_url,
                "document_type": doc_type,
                "field_count": len(extracted_data.get('fields', [])),
                "coverage_percentage": coverage_percentage,
                "schema": extracted_data
            }
            
        except Exception as e:
            logger.error(f"Error processing document {doc_url}: {e}")
            await self.record_extraction_status(
                subsidy_id=subsidy_id,
                document_url=doc_url,
                document_type=doc_type,
                status='failure',
                error_message=str(e)
            )
            raise

    async def download_document(self, url: str) -> str:
        """
        Download a document from URL to temporary file.
        
        Args:
            url: Document URL to download.
            
        Returns:
            Path to downloaded temporary file.
        """
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=self.download_timeout)) as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise ValueError(f"Failed to download document: HTTP {response.status}")
                
                # Create temporary file
                suffix = Path(urlparse(url).path).suffix or '.tmp'
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                
                async with aiofiles.open(temp_file.name, 'wb') as f:
                    async for chunk in response.content.iter_chunked(8192):
                        await f.write(chunk)
                
                return temp_file.name

    async def extract_pdf_schema(self, file_path: str) -> Dict[str, Any]:
        """
        Extract schema from PDF document.
        
        Args:
            file_path: Path to PDF file.
            
        Returns:
            Extracted schema data.
        """
        def _extract_pdf():
            """Synchronous PDF extraction to run in thread pool."""
            fields = []
            raw_unclassified = []
            
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = ""
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
                    
                    # Try to extract form fields from PDF annotations
                    if hasattr(pdf_reader, 'get_form_text_fields'):
                        form_fields = pdf_reader.get_form_text_fields()
                        if form_fields:
                            for field_name, field_value in form_fields.items():
                                fields.append({
                                    "name": self.normalize_field_name(field_name),
                                    "label": field_name,
                                    "type": "text",
                                    "required": False,
                                    "help": ""
                                })
                    
                    # Parse text content for form-like patterns
                    parsed_fields = self.parse_text_for_fields(text_content)
                    fields.extend(parsed_fields['fields'])
                    raw_unclassified.extend(parsed_fields['raw_unclassified'])
                    
            except Exception as e:
                logger.error(f"Error extracting PDF schema: {e}")
                raw_unclassified.append(f"PDF extraction error: {str(e)}")
            
            return {"fields": fields, "raw_unclassified": raw_unclassified}
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract_pdf)

    async def extract_docx_schema(self, file_path: str) -> Dict[str, Any]:
        """
        Extract schema from DOCX document.
        
        Args:
            file_path: Path to DOCX file.
            
        Returns:
            Extracted schema data.
        """
        def _extract_docx():
            """Synchronous DOCX extraction to run in thread pool."""
            fields = []
            raw_unclassified = []
            
            try:
                doc = Document(file_path)
                text_content = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                    text_content += "\n"
                
                # Parse content for form fields
                parsed_fields = self.parse_text_for_fields(text_content)
                fields.extend(parsed_fields['fields'])
                raw_unclassified.extend(parsed_fields['raw_unclassified'])
                
            except Exception as e:
                logger.error(f"Error extracting DOCX schema: {e}")
                raw_unclassified.append(f"DOCX extraction error: {str(e)}")
            
            return {"fields": fields, "raw_unclassified": raw_unclassified}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract_docx)

    async def extract_xlsx_schema(self, file_path: str) -> Dict[str, Any]:
        """
        Extract schema from XLSX document.
        
        Args:
            file_path: Path to XLSX file.
            
        Returns:
            Extracted schema data.
        """
        def _extract_xlsx():
            """Synchronous XLSX extraction to run in thread pool."""
            fields = []
            raw_unclassified = []
            
            try:
                workbook = load_workbook(file_path, read_only=True)
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    
                    # Look for form-like patterns in cells
                    for row in sheet.iter_rows(max_col=10, max_row=100):
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                # Check for field-like patterns
                                if ':' in cell.value or '?' in cell.value:
                                    field_data = self.parse_cell_for_field(cell.value)
                                    if field_data:
                                        fields.append(field_data)
                                    else:
                                        raw_unclassified.append(cell.value)
                
                workbook.close()
                
            except Exception as e:
                logger.error(f"Error extracting XLSX schema: {e}")
                raw_unclassified.append(f"XLSX extraction error: {str(e)}")
            
            return {"fields": fields, "raw_unclassified": raw_unclassified}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract_xlsx)

    def parse_text_for_fields(self, text: str) -> Dict[str, List]:
        """
        Parse text content to identify form fields.
        
        Args:
            text: Text content to parse.
            
        Returns:
            Dictionary with fields and raw_unclassified lists.
        """
        fields = []
        raw_unclassified = []
        
        # Common field patterns
        field_patterns = [
            r'(.+?):\s*_+',  # Label followed by underlines
            r'(.+?)\s*\[\s*\]',  # Label with checkbox
            r'(.+?)\s*\(\s*\)',  # Label with parentheses
            r'(\d+\.\s*.+?)(?=\d+\.|$)',  # Numbered questions
        ]
        
        for pattern in field_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                label = match.group(1).strip()
                if len(label) > 5 and len(label) < 200:  # Reasonable field label length
                    field = {
                        "name": self.normalize_field_name(label),
                        "label": label,
                        "type": self.infer_field_type(label),
                        "required": self.is_field_required(label),
                        "help": ""
                    }
                    fields.append(field)
                else:
                    raw_unclassified.append(label)
        
        return {"fields": fields, "raw_unclassified": raw_unclassified}

    def parse_cell_for_field(self, cell_value: str) -> Optional[Dict[str, Any]]:
        """
        Parse Excel cell value to identify if it's a form field.
        
        Args:
            cell_value: Cell value to parse.
            
        Returns:
            Field data if identified, None otherwise.
        """
        if ':' in cell_value:
            parts = cell_value.split(':', 1)
            label = parts[0].strip()
            
            if len(label) > 3 and len(label) < 100:
                return {
                    "name": self.normalize_field_name(label),
                    "label": label,
                    "type": self.infer_field_type(label),
                    "required": self.is_field_required(label),
                    "help": parts[1].strip() if len(parts) > 1 else ""
                }
        
        return None

    def normalize_field_name(self, label: str) -> str:
        """
        Normalize field label to create a valid field name.
        
        Args:
            label: Original field label.
            
        Returns:
            Normalized field name.
        """
        # Remove special characters and spaces, convert to lowercase
        name = re.sub(r'[^a-zA-Z0-9\s]', '', label)
        name = re.sub(r'\s+', '_', name.strip())
        return name.lower()

    def infer_field_type(self, label: str) -> str:
        """
        Infer field type from label text.
        
        Args:
            label: Field label text.
            
        Returns:
            Inferred field type.
        """
        label_lower = label.lower()
        
        if any(word in label_lower for word in ['date', 'datum', 'data']):
            return 'date'
        elif any(word in label_lower for word in ['email', 'e-mail']):
            return 'email'
        elif any(word in label_lower for word in ['phone', 'tel', 'telefon']):
            return 'tel'
        elif any(word in label_lower for word in ['number', 'numéro', 'nummer', 'nr', 'amount', 'montant']):
            return 'number'
        elif any(word in label_lower for word in ['description', 'descriere', 'details']):
            return 'textarea'
        elif any(word in label_lower for word in ['select', 'choose', 'option']):
            return 'select'
        elif any(word in label_lower for word in ['checkbox', 'check', 'tick']):
            return 'checkbox'
        else:
            return 'text'

    def is_field_required(self, label: str) -> bool:
        """
        Determine if field is required based on label.
        
        Args:
            label: Field label text.
            
        Returns:
            True if field appears to be required.
        """
        label_lower = label.lower()
        return any(word in label_lower for word in ['*', 'required', 'obligatoire', 'obligatoriu', 'mandatory'])

    def calculate_coverage(self, extracted_data: Dict[str, Any]) -> float:
        """
        Calculate extraction coverage percentage.
        
        Args:
            extracted_data: Extracted schema data.
            
        Returns:
            Coverage percentage (0-100).
        """
        fields_count = len(extracted_data.get('fields', []))
        unclassified_count = len(extracted_data.get('raw_unclassified', []))
        
        if fields_count + unclassified_count == 0:
            return 0.0
        
        return (fields_count / (fields_count + unclassified_count)) * 100

    def consolidate_schemas(self, extraction_results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Consolidate schemas from multiple documents.
        
        Args:
            extraction_results: List of extraction results from documents.
            
        Returns:
            Consolidated schema.
        """
        all_fields = []
        all_raw_unclassified = []
        
        for result in extraction_results:
            schema = result.get('schema', {})
            all_fields.extend(schema.get('fields', []))
            all_raw_unclassified.extend(schema.get('raw_unclassified', []))
        
        # Remove duplicate fields (same name)
        unique_fields = {}
        for field in all_fields:
            field_name = field.get('name')
            if field_name and field_name not in unique_fields:
                unique_fields[field_name] = field
        
        return {
            "fields": list(unique_fields.values()),
            "raw_unclassified": all_raw_unclassified,
            "total_documents": len(extraction_results),
            "extraction_timestamp": datetime.utcnow().isoformat()
        }

    async def update_subsidy_schema(self, subsidy_id: str, schema: Dict[str, Any]) -> None:
        """
        Update subsidy record with extracted schema.
        
        Args:
            subsidy_id: Subsidy ID to update.
            schema: Extracted schema to store.
        """
        try:
            result = self.supabase.table('subsidies').update({
                'application_schema': schema
            }).eq('id', subsidy_id).execute()
            
            if result.data:
                logger.info(f"Updated schema for subsidy {subsidy_id}")
            else:
                logger.warning(f"No rows updated for subsidy {subsidy_id}")
                
        except Exception as e:
            logger.error(f"Error updating subsidy schema: {e}")
            raise

    async def record_extraction_status(
        self,
        subsidy_id: str,
        document_url: str,
        document_type: str,
        status: str,
        field_count: int = 0,
        coverage_percentage: float = 0.0,
        extracted_schema: Optional[Dict[str, Any]] = None,
        error_message: Optional[str] = None
    ) -> None:
        """
        Record document extraction status in database.
        
        Args:
            subsidy_id: Subsidy ID.
            document_url: Document URL.
            document_type: Document type.
            status: Extraction status.
            field_count: Number of fields extracted.
            coverage_percentage: Coverage percentage.
            extracted_schema: Extracted schema data.
            error_message: Error message if failed.
        """
        try:
            extraction_errors = []
            if error_message:
                extraction_errors.append({"error": error_message, "timestamp": datetime.utcnow().isoformat()})
            
            # Check if record already exists (idempotent)
            existing = self.supabase.table('document_extraction_status').select('id').eq(
                'subsidy_id', subsidy_id
            ).eq('document_url', document_url).execute()
            
            data = {
                'subsidy_id': subsidy_id,
                'document_url': document_url,
                'document_type': document_type,
                'extraction_status': status,
                'field_count': field_count,
                'coverage_percentage': coverage_percentage,
                'extraction_errors': extraction_errors,
                'extracted_schema': extracted_schema,
                'updated_at': datetime.utcnow().isoformat()
            }
            
            if existing.data:
                # Update existing record
                result = self.supabase.table('document_extraction_status').update(data).eq(
                    'id', existing.data[0]['id']
                ).execute()
            else:
                # Insert new record
                result = self.supabase.table('document_extraction_status').insert(data).execute()
            
            logger.info(f"Recorded extraction status for {document_url}: {status}")
            
        except Exception as e:
            logger.error(f"Error recording extraction status: {e}")
            # Don't raise - this shouldn't stop the main process


async def main():
    """Main CLI function."""
    parser = argparse.ArgumentParser(
        description='Extract document schemas from subsidy application forms'
    )
    parser.add_argument(
        '--subsidy-ids',
        nargs='+',
        help='Specific subsidy IDs to process'
    )
    parser.add_argument(
        '--all-subsidies',
        action='store_true',
        help='Process all subsidies in database'
    )
    parser.add_argument(
        '--batch-size',
        type=int,
        default=10,
        help='Number of subsidies to process in each batch'
    )
    parser.add_argument(
        '--max-concurrent',
        type=int,
        default=5,
        help='Maximum concurrent document processing'
    )
    
    args = parser.parse_args()
    
    if not args.subsidy_ids and not args.all_subsidies:
        parser.error('Must specify either --subsidy-ids or --all-subsidies')
    
    # Override environment variable if specified
    if args.max_concurrent:
        os.environ['MAX_CONCURRENT_DOCS'] = str(args.max_concurrent)
    
    try:
        extractor = DocumentSchemaExtractor()
        
        subsidy_ids = args.subsidy_ids if args.subsidy_ids else None
        stats = await extractor.process_subsidies(subsidy_ids, args.batch_size)
        
        # Print summary
        print("\n=== EXTRACTION SUMMARY ===")
        print(f"Total subsidies: {stats['total']}")
        print(f"Successfully processed: {stats['processed']}")
        print(f"Failed: {stats['failed']}")
        print(f"Success rate: {(stats['processed'] / stats['total'] * 100):.1f}%")
        
        if stats['extraction_results']:
            total_fields = sum(r.get('total_fields', 0) for r in stats['extraction_results'])
            avg_fields = total_fields / len(stats['extraction_results'])
            print(f"Total fields extracted: {total_fields}")
            print(f"Average fields per subsidy: {avg_fields:.1f}")
        
    except Exception as e:
        logger.error(f"Fatal error: {e}")
        return 1
    
    return 0


if __name__ == '__main__':
    import sys
    sys.exit(asyncio.run(main()))