#!/usr/bin/env python3
"""
Python Document Extraction Pipeline
===================================

Pure Python document extraction system to replace Tika completely.
Supports PDF, DOCX, XLS/XLSX, ODT, images, and fallback to unstructured library.
Features OCR for scanned documents and comprehensive error handling.
"""

import os
import io
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
import hashlib

# PDF extraction
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.pdfpage import PDFTextExtractionNotAllowed
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

# DOCX extraction
try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Excel extraction
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# ODT extraction
try:
    from odf.opendocument import load as odf_load
    from odf.text import P as OdfParagraph
    from odf.teletype import extractText as odf_extract_text
    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False

# OCR capabilities
try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

# Fallback extraction
try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False

# Additional libraries
import subprocess
import requests
from datetime import datetime


class DocumentExtractionResult:
    """Container for document extraction results"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        self.file_type = self._detect_file_type()
        self.extraction_method = None
        self.text_content = ""
        self.metadata = {}
        self.success = False
        self.error = None
        self.warnings = []
        self.extraction_time = 0
        self.character_count = 0
        self.page_count = None
        self.confidence_score = 0.0
        
    def _detect_file_type(self) -> str:
        """Detect file type from extension"""
        ext = Path(self.file_path).suffix.lower()
        return ext[1:] if ext else 'unknown'
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert result to dictionary"""
        return {
            'file_path': self.file_path,
            'file_name': self.file_name,
            'file_size': self.file_size,
            'file_type': self.file_type,
            'extraction_method': self.extraction_method,
            'text_content': self.text_content,
            'metadata': self.metadata,
            'success': self.success,
            'error': self.error,
            'warnings': self.warnings,
            'extraction_time': self.extraction_time,
            'character_count': self.character_count,
            'page_count': self.page_count,
            'confidence_score': self.confidence_score
        }


class PythonDocumentExtractor:
    """Pure Python document extraction system"""
    
    def __init__(
        self,
        enable_ocr: bool = True,
        ocr_language: str = 'eng+fra+ron',  # English, French, Romanian
        max_file_size_mb: float = 50.0,
        temp_dir: Optional[str] = None
    ):
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language
        self.max_file_size_bytes = int(max_file_size_mb * 1024 * 1024)
        self.temp_dir = temp_dir or tempfile.gettempdir()
        
        # Setup logging
        self.logger = logging.getLogger(__name__)
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
            self.logger.setLevel(logging.INFO)
            
        self._check_dependencies()
        
    def _check_dependencies(self):
        """Check and log available dependencies"""
        deps = {
            'pdfplumber': PDFPLUMBER_AVAILABLE,
            'pdfminer': PDFMINER_AVAILABLE,
            'python-docx': PYTHON_DOCX_AVAILABLE,
            'openpyxl': OPENPYXL_AVAILABLE,
            'pandas': PANDAS_AVAILABLE,
            'odfpy': ODF_AVAILABLE,
            'pytesseract': PYTESSERACT_AVAILABLE,
            'unstructured': UNSTRUCTURED_AVAILABLE
        }
        
        available = [name for name, avail in deps.items() if avail]
        missing = [name for name, avail in deps.items() if not avail]
        
        self.logger.info(f"📚 Available libraries: {', '.join(available)}")
        if missing:
            self.logger.warning(f"⚠️ Missing libraries: {', '.join(missing)}")
            
    def extract_document(self, file_path: str) -> DocumentExtractionResult:
        """Main document extraction method"""
        start_time = datetime.now()
        result = DocumentExtractionResult(file_path)
        
        try:
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
                
            if result.file_size > self.max_file_size_bytes:
                raise ValueError(f"File too large: {result.file_size / 1024 / 1024:.1f}MB > {self.max_file_size_bytes / 1024 / 1024:.1f}MB")
            
            self.logger.info(f"📄 Extracting text from: {result.file_name} ({result.file_type})")
            
            # Route to appropriate extraction method
            if result.file_type in ['pdf']:
                self._extract_pdf(result)
            elif result.file_type in ['docx', 'doc']:
                self._extract_docx(result)
            elif result.file_type in ['xlsx', 'xls']:
                self._extract_excel(result)
            elif result.file_type in ['odt']:
                self._extract_odt(result)
            elif result.file_type in ['txt', 'csv']:
                self._extract_plain_text(result)
            elif result.file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                self._extract_image_text(result)
            else:
                # Fallback to unstructured
                self._extract_with_unstructured(result)
                
            # Post-processing
            self._post_process_text(result)
            result.character_count = len(result.text_content)
            result.success = True
            
        except Exception as e:
            result.error = str(e)
            result.success = False
            self.logger.error(f"❌ Extraction failed for {result.file_name}: {e}")
            
        finally:
            result.extraction_time = (datetime.now() - start_time).total_seconds()
            
        self.logger.info(f"✅ Extraction complete: {result.character_count} chars, {result.extraction_time:.2f}s")
        return result
        
    def _extract_pdf(self, result: DocumentExtractionResult):
        """Extract text from PDF files"""
        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                result.extraction_method = 'pdfplumber'
                text_parts = []
                
                with pdfplumber.open(result.file_path) as pdf:
                    result.page_count = len(pdf.pages)
                    result.metadata['pdf_info'] = pdf.metadata
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        else:
                            result.warnings.append(f"No text found on page {page_num + 1}")
                            
                result.text_content = '\n\n'.join(text_parts)
                result.confidence_score = 0.9
                
                # If no text extracted, try OCR
                if not result.text_content.strip() and self.enable_ocr:
                    result.warnings.append("No text extracted, attempting OCR")
                    self._extract_pdf_with_ocr(result)
                    
                return
                
            except Exception as e:
                result.warnings.append(f"pdfplumber failed: {e}")
        
        # Fallback to pdfminer
        if PDFMINER_AVAILABLE:
            try:
                result.extraction_method = 'pdfminer'
                result.text_content = pdfminer_extract_text(result.file_path)
                result.confidence_score = 0.8
                
                # Count pages
                with open(result.file_path, 'rb') as file:
                    result.page_count = len(list(PDFPage.get_pages(file)))
                    
                if not result.text_content.strip() and self.enable_ocr:
                    result.warnings.append("No text extracted with pdfminer, attempting OCR")
                    self._extract_pdf_with_ocr(result)
                    
                return
                
            except Exception as e:
                result.warnings.append(f"pdfminer failed: {e}")
        
        # Last resort: OCR if enabled
        if self.enable_ocr:
            self._extract_pdf_with_ocr(result)
        else:
            raise RuntimeError("No PDF extraction libraries available and OCR disabled")
            
    def _extract_pdf_with_ocr(self, result: DocumentExtractionResult):
        """Extract text from PDF using OCR"""
        if not PYTESSERACT_AVAILABLE:
            raise RuntimeError("pytesseract not available for OCR")
            
        try:
            result.extraction_method = 'pytesseract_ocr'
            
            # Convert PDF to images and OCR each page
            import fitz  # PyMuPDF for PDF to image conversion
            
            doc = fitz.open(result.file_path)
            result.page_count = len(doc)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # OCR the image
                image = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(
                    image, 
                    lang=self.ocr_language,
                    config='--psm 6'  # Uniform text block
                )
                
                if page_text.strip():
                    text_parts.append(page_text)
                    
            doc.close()
            
            result.text_content = '\n\n'.join(text_parts)
            result.confidence_score = 0.7  # OCR is less reliable
            result.warnings.append("Text extracted using OCR")
            
        except Exception as e:
            # Try alternative OCR method without PyMuPDF
            try:
                self._extract_pdf_ocr_alternative(result)
            except Exception as e2:
                raise RuntimeError(f"OCR extraction failed: {e}, {e2}")
                
    def _extract_pdf_ocr_alternative(self, result: DocumentExtractionResult):
        """Alternative OCR method using pdf2image"""
        try:
            from pdf2image import convert_from_path
            
            result.extraction_method = 'pdf2image_ocr'
            
            # Convert PDF to images
            images = convert_from_path(result.file_path)
            result.page_count = len(images)
            text_parts = []
            
            for page_num, image in enumerate(images):
                page_text = pytesseract.image_to_string(
                    image, 
                    lang=self.ocr_language,
                    config='--psm 6'
                )
                
                if page_text.strip():
                    text_parts.append(page_text)
                    
            result.text_content = '\n\n'.join(text_parts)
            result.confidence_score = 0.7
            result.warnings.append("Text extracted using pdf2image + OCR")
            
        except Exception as e:
            raise RuntimeError(f"Alternative OCR failed: {e}")
            
    def _extract_docx(self, result: DocumentExtractionResult):
        """Extract text from DOCX files"""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
            
        try:
            result.extraction_method = 'python-docx'
            
            doc = DocxDocument(result.file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append('\t'.join(row_text))
                        
            result.text_content = '\n'.join(text_parts)
            result.confidence_score = 0.95
            result.metadata['docx_paragraphs'] = len(doc.paragraphs)
            result.metadata['docx_tables'] = len(doc.tables)
            
        except Exception as e:
            if UNSTRUCTURED_AVAILABLE:
                result.warnings.append(f"python-docx failed: {e}, trying unstructured")
                self._extract_with_unstructured(result)
            else:
                raise RuntimeError(f"DOCX extraction failed: {e}")
                
    def _extract_excel(self, result: DocumentExtractionResult):
        """Extract text from Excel files"""
        # Try openpyxl first
        if OPENPYXL_AVAILABLE:
            try:
                result.extraction_method = 'openpyxl'
                
                workbook = openpyxl.load_workbook(result.file_path, data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            text_parts.append('\t'.join(row_text))
                            
                result.text_content = '\n'.join(text_parts)
                result.confidence_score = 0.9
                result.metadata['excel_sheets'] = len(workbook.sheetnames)
                
                return
                
            except Exception as e:
                result.warnings.append(f"openpyxl failed: {e}")
                
        # Fallback to pandas
        if PANDAS_AVAILABLE:
            try:
                result.extraction_method = 'pandas'
                
                # Read all sheets
                excel_file = pd.ExcelFile(result.file_path)
                text_parts = []
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    text_parts.append(df.to_string(index=False))
                    
                result.text_content = '\n\n'.join(text_parts)
                result.confidence_score = 0.85
                result.metadata['excel_sheets'] = len(excel_file.sheet_names)
                
                return
                
            except Exception as e:
                result.warnings.append(f"pandas failed: {e}")
                
        raise RuntimeError("No Excel extraction libraries available")
        
    def _extract_odt(self, result: DocumentExtractionResult):
        """Extract text from ODT files"""
        if not ODF_AVAILABLE:
            raise RuntimeError("odfpy not available")
            
        try:
            result.extraction_method = 'odfpy'
            
            doc = odf_load(result.file_path)
            text_parts = []
            
            # Extract paragraphs
            paragraphs = doc.getElementsByType(OdfParagraph)
            for para in paragraphs:
                text = odf_extract_text(para)
                if text.strip():
                    text_parts.append(text)
                    
            result.text_content = '\n'.join(text_parts)
            result.confidence_score = 0.9
            result.metadata['odt_paragraphs'] = len(paragraphs)
            
        except Exception as e:
            if UNSTRUCTURED_AVAILABLE:
                result.warnings.append(f"odfpy failed: {e}, trying unstructured")
                self._extract_with_unstructured(result)
            else:
                raise RuntimeError(f"ODT extraction failed: {e}")
                
    def _extract_plain_text(self, result: DocumentExtractionResult):
        """Extract text from plain text files"""
        try:
            result.extraction_method = 'plain_text'
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(result.file_path, 'r', encoding=encoding) as file:
                        result.text_content = file.read()
                    result.confidence_score = 1.0
                    result.metadata['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise RuntimeError("Could not decode file with any encoding")
                
        except Exception as e:
            raise RuntimeError(f"Plain text extraction failed: {e}")
            
    def _extract_image_text(self, result: DocumentExtractionResult):
        """Extract text from images using OCR"""
        if not PYTESSERACT_AVAILABLE:
            raise RuntimeError("pytesseract not available for image OCR")
            
        try:
            result.extraction_method = 'pytesseract_image'
            
            image = Image.open(result.file_path)
            result.text_content = pytesseract.image_to_string(
                image, 
                lang=self.ocr_language,
                config='--psm 6'
            )
            
            result.confidence_score = 0.8
            result.metadata['image_size'] = image.size
            result.metadata['image_mode'] = image.mode
            
        except Exception as e:
            raise RuntimeError(f"Image OCR failed: {e}")
            
    def _extract_with_unstructured(self, result: DocumentExtractionResult):
        """Fallback extraction using unstructured library"""
        if not UNSTRUCTURED_AVAILABLE:
            raise RuntimeError("unstructured library not available")
            
        try:
            result.extraction_method = 'unstructured'
            
            elements = partition(filename=result.file_path)
            text_parts = []
            
            for element in elements:
                if hasattr(element, 'text') and element.text.strip():
                    text_parts.append(element.text)
                    
            result.text_content = '\n'.join(text_parts)
            result.confidence_score = 0.7
            result.metadata['unstructured_elements'] = len(elements)
            
        except Exception as e:
            raise RuntimeError(f"Unstructured extraction failed: {e}")
            
    def _post_process_text(self, result: DocumentExtractionResult):
        """Clean and normalize extracted text"""
        if not result.text_content:
            return
            
        text = result.text_content
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
                
        # Join with single newlines
        result.text_content = '\n'.join(cleaned_lines)
        
        # Remove very long sequences of repeated characters
        import re
        result.text_content = re.sub(r'(.)\1{20,}', r'\1\1\1', result.text_content)
        
    def extract_multiple_documents(self, file_paths: List[str]) -> Dict[str, DocumentExtractionResult]:
        """Extract text from multiple documents"""
        results = {}
        
        for file_path in file_paths:
            try:
                result = self.extract_document(file_path)
                results[file_path] = result
            except Exception as e:
                self.logger.error(f"❌ Failed to extract {file_path}: {e}")
                results[file_path] = DocumentExtractionResult(file_path)
                results[file_path].error = str(e)
                
        return results
    
    def extract_document_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from document and return as dictionary.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dict containing extraction results
        """
        result = self.extract_document(file_path)
        return result.to_dict()


def extract_document_text(
    file_path: str,
    enable_ocr: bool = True,
    ocr_language: str = 'eng+fra+ron',
    max_file_size_mb: float = 50.0
) -> Dict[str, Any]:
    """
    Convenience function for document text extraction.
    
    Args:
        file_path: Path to document file
        enable_ocr: Enable OCR for scanned documents
        ocr_language: OCR language codes
        max_file_size_mb: Maximum file size in MB
        
    Returns:
        Dict containing extraction results
    """
    extractor = PythonDocumentExtractor(
        enable_ocr=enable_ocr,
        ocr_language=ocr_language,
        max_file_size_mb=max_file_size_mb
    )
    
    result = extractor.extract_document(file_path)
    return result.to_dict()


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python python_document_extractor.py <document_file>")
        sys.exit(1)
    
    document_file = sys.argv[1]
    
    if not os.path.exists(document_file):
        print(f"File not found: {document_file}")
        sys.exit(1)
    
    try:
        result = extract_document_text(document_file)
        
        print(f"\n{'='*60}")
        print("DOCUMENT EXTRACTION RESULTS")
        print('='*60)
        print(f"File: {result['file_name']}")
        print(f"Type: {result['file_type']}")
        print(f"Method: {result['extraction_method']}")
        print(f"Success: {result['success']}")
        print(f"Characters: {result['character_count']}")
        print(f"Time: {result['extraction_time']:.2f}s")
        
        if result['warnings']:
            print(f"Warnings: {', '.join(result['warnings'])}")
            
        if result['error']:
            print(f"Error: {result['error']}")
        else:
            print(f"\n--- First 500 characters ---")
            print(result['text_content'][:500] + "..." if len(result['text_content']) > 500 else result['text_content'])
            
        print('='*60)
        
    except Exception as e:
        print(f"❌ Extraction failed: {e}")
        sys.exit(1)